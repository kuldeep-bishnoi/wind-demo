import os
import io
import base64
import tempfile
import logging
import mimetypes
from logging.handlers import RotatingFileHandler
from typing import Dict, Any, Optional, Union, List, Tuple

import streamlit as st
from fpdf import FPDF
from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
import magic
from PIL import Image
import openai
import yaml

# Set up logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
handler = RotatingFileHandler('app.log', maxBytes=100000, backupCount=3)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)

# Set page config at the top
st.set_page_config(
    page_title="Universal File Converter",
    page_icon="📁",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Supported file types and their handlers
SUPPORTED_TYPES = {
    'Documents': ['pdf', 'docx', 'txt', 'md'],
    'Spreadsheets': ['csv', 'xlsx', 'xls'],
    'Images': ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'],
    'Data': ['json', 'yaml', 'yml', 'xml'],
    'Code': ['py', 'js', 'html', 'css', 'java', 'c', 'cpp', 'cs', 'go', 'php', 'rb', 'swift']
}

# File type icons
FILE_ICONS = {
    'pdf': '📄', 'docx': '📝', 'txt': '📝', 'md': '📝',
    'csv': '📊', 'xlsx': '📊', 'xls': '📊',
    'png': '🖼️', 'jpg': '🖼️', 'jpeg': '🖼️', 'gif': '🖼️', 'bmp': '🖼️', 'tiff': '🖼️', 'webp': '🖼️',
    'json': '📋', 'yaml': '📋', 'yml': '📋', 'xml': '📋',
    'py': '🐍', 'js': '📜', 'html': '🌐', 'css': '🎨', 'java': '☕', 'c': '🔧', 'cpp': '🔧', 'cs': '🔷', 'go': '🐹', 'php': '🐘', 'rb': '💎', 'swift': '🐦'
}

def detect_file_type(file_data: bytes, filename: str) -> str:
    """Detect file type using both magic and extension."""
    try:
        # First try magic for content-based detection
        mime = magic.Magic(mime=True)
        file_type = mime.from_buffer(file_data)
        
        # Extract extension from filename as fallback
        _, ext = os.path.splitext(filename)
        ext = ext.lower().lstrip('.')
        
        # If magic couldn't determine or gave generic types
        if not file_type or file_type == 'application/octet-stream':
            return ext if ext in [e for types in SUPPORTED_TYPES.values() for e in types] else 'unknown'
            
        # Map MIME types to extensions
        mime_to_ext = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/markdown': 'md',
            'text/csv': 'csv',
            'application/vnd.ms-excel': 'xls',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'image/png': 'png',
            'image/jpeg': 'jpg',
            'image/gif': 'gif',
            'image/bmp': 'bmp',
            'image/tiff': 'tiff',
            'image/webp': 'webp',
            'application/json': 'json',
            'text/yaml': 'yaml',
            'application/xml': 'xml',
            'text/x-python': 'py',
            'application/javascript': 'js',
            'text/html': 'html',
            'text/css': 'css',
            'text/x-java-source': 'java',
            'text/x-c': 'c',
            'text/x-c++': 'cpp',
            'text/x-csharp': 'cs',
            'text/x-go': 'go',
            'application/x-httpd-php': 'php',
            'application/x-ruby': 'rb',
            'text/x-swift': 'swift'
        }
        
        return mime_to_ext.get(file_type.split(';')[0], ext if ext else 'unknown')
    except Exception as e:
        logger.error(f"Error detecting file type: {e}")
        _, ext = os.path.splitext(filename)
        return ext.lower().lstrip('.') if ext else 'unknown'

def read_file_content(file_data: bytes, file_type: str) -> Tuple[Any, str]:
    """Read content from different file types."""
    try:
        if file_type in ['pdf']:
            pdf = PdfReader(io.BytesIO(file_data))
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n\n"
            return text, "text/plain"
            
        elif file_type in ['docx']:
            doc = Document(io.BytesIO(file_data))
            return "\n".join([para.text for para in doc.paragraphs]), "text/plain"
            
        elif file_type in ['csv']:
            df = pd.read_csv(io.BytesIO(file_data))
            return df, "dataframe"
            
        elif file_type in ['xlsx', 'xls']:
            df = pd.read_excel(io.BytesIO(file_data))
            return df, "dataframe"
            
        elif file_type in ['json']:
            data = pd.read_json(io.BytesIO(file_data))
            if isinstance(data, dict):
                return data, "dict"
            return data, "dataframe"
            
        elif file_type in ['yaml', 'yml']:
            data = yaml.safe_load(io.BytesIO(file_data))
            if isinstance(data, dict):
                return data, "dict"
            return str(data), "text/plain"
            
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
            return file_data, f"image/{file_type}"
            
        else:  # For text files and code files
            try:
                return file_data.decode('utf-8'), "text/plain"
            except UnicodeDecodeError:
                return file_data, "application/octet-stream"
                
    except Exception as e:
        logger.error(f"Error reading {file_type} file: {e}")
        raise

def generate_summary(text: str, api_key: str) -> Optional[str]:
    """Generate a summary using OpenAI's API."""
    try:
        openai.api_key = api_key
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that summarizes text concisely."},
                {"role": "user", "content": f"Please provide a concise summary of the following text:\n\n{text[:3000]}"}
            ],
            max_tokens=300,
            temperature=0.5
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        logger.error(f"Error generating summary: {e}")
        return None

def display_file_preview(content, content_type: str, file_type: str):
    """Display file preview based on content type."""
    st.subheader("File Preview")
    
    if content_type == "dataframe":
        st.dataframe(content)
    elif content_type.startswith("image"):
        st.image(content, use_column_width=True)
    elif content_type == "dict":
        st.json(content)
    else:  # text/plain
        st.text_area("File Content", value=content, height=300)

def convert_file(content, target_format: str, original_type: str):
    """Convert file content to target format."""
    if target_format == original_type:
        return content  # No conversion needed
        
    try:
        if target_format == 'pdf':
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            if isinstance(content, str):
                pdf.multi_cell(0, 10, txt=content)
            elif isinstance(content, dict):
                for key, value in content.items():
                    pdf.cell(0, 10, txt=f"{key}: {value}", ln=True)
            elif hasattr(content, 'to_string'):  # For DataFrames
                pdf.multi_cell(0, 10, txt=content.to_string())
                
            output = io.BytesIO()
            pdf_output = pdf.output(dest='S').encode('latin1')
            output.write(pdf_output)
            return output.getvalue(), "application/pdf"
            
        elif target_format == 'docx':
            doc = Document()
            
            if isinstance(content, str):
                doc.add_paragraph(content)
            elif isinstance(content, dict):
                for key, value in content.items():
                    doc.add_paragraph(f"{key}: {value}")
            elif hasattr(content, 'to_string'):  # For DataFrames
                doc.add_paragraph(content.to_string())
                
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif target_format in ['csv', 'xlsx']:
            if isinstance(content, (dict, list)):
                df = pd.DataFrame(content)
            elif hasattr(content, 'to_dict'):
                df = pd.DataFrame(content.to_dict())
            else:
                df = pd.DataFrame({"Content": [content]})
                
            output = io.BytesIO()
            if target_format == 'csv':
                df.to_csv(output, index=False)
                return output.getvalue(), "text/csv"
            else:
                df.to_excel(output, index=False)
                return output.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                
        elif target_format in ['json', 'yaml', 'yml']:
            if hasattr(content, 'to_dict'):
                data = content.to_dict()
            elif isinstance(content, str):
                data = {"content": content}
            else:
                data = content
                
            output = io.BytesIO()
            if target_format == 'json':
                json_str = json.dumps(data, indent=2)
                output.write(json_str.encode('utf-8'))
                return output.getvalue(), "application/json"
            else:
                yaml_str = yaml.dump(data, default_flow_style=False)
                output.write(yaml_str.encode('utf-8'))
                return output.getvalue(), "text/yaml"
                
    except Exception as e:
        logger.error(f"Error converting to {target_format}: {e}")
        raise
    
    return content, f"application/{target_format}"

def main():
    """Main application function."""
    st.title("📁 Universal File Converter")
    st.markdown("Upload files to preview, convert, and export in different formats.")
    
    # Sidebar for settings
    with st.sidebar:
        st.header("Settings")
        api_key = st.text_input("OpenAI API Key (for summaries)", type="password", 
                              help="Optional: Enter your OpenAI API key to enable AI-powered summaries.")
        
        st.markdown("---")
        st.markdown("### Supported File Types")
        for category, extensions in SUPPORTED_TYPES.items():
            st.markdown(f"**{category}**: {', '.join(extensions)}")
        
        st.markdown("---")
        st.markdown("### About")
        st.markdown("This tool allows you to preview, convert, and export files in various formats.")
    
    # File upload section
    uploaded_files = st.file_uploader(
        "Upload files", 
        type=[ext for exts in SUPPORTED_TYPES.values() for ext in exts], 
        accept_multiple_files=True,
        help="Upload one or more files to convert"
    )
    
    if not uploaded_files:
        st.info("Please upload one or more files to get started.")
        return
    
    # Process each uploaded file
    for file in uploaded_files:
        file_data = file.getvalue()
        file_type = detect_file_type(file_data, file.name)
        
        with st.expander(f"{FILE_ICONS.get(file_type, '📄')} {file.name}", expanded=True):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                try:
                    content, content_type = read_file_content(file_data, file_type)
                    display_file_preview(content, content_type, file_type)
                except Exception as e:
                    st.error(f"Error reading file: {str(e)}")
                    continue
            
            with col2:
                st.subheader("Actions")
                
                # File info
                file_size = len(file_data) / 1024  # KB
                st.caption(f"Type: {file_type.upper()}")
                st.caption(f"Size: {file_size:.2f} KB")
                
                # Generate summary if API key is provided
                if api_key and content_type == "text/plain" and len(str(content)) > 100:
                    if st.button("Generate Summary", key=f"summarize_{file.name}"):
                        with st.spinner("Generating summary..."):
                            summary = generate_summary(str(content), api_key)
                            if summary:
                                st.subheader("AI Summary")
                                st.write(summary)
                
                # Conversion options
                st.markdown("### Convert To")
                target_format = st.selectbox(
                    "Select target format",
                    ["pdf", "docx", "txt", "csv", "xlsx", "json", "yaml"],
                    key=f"format_{file.name}",
                    format_func=lambda x: x.upper()
                )
                
                # Convert and download
                if st.button("Convert & Download", key=f"convert_{file.name}"):
                    with st.spinner(f"Converting to {target_format.upper()}..."):
                        try:
                            output_data, mime_type = convert_file(content, target_format, file_type)
                            
                            # Determine file extension
                            ext_map = {
                                'pdf': 'pdf',
                                'docx': 'docx',
                                'txt': 'txt',
                                'csv': 'csv',
                                'xlsx': 'xlsx',
                                'json': 'json',
                                'yaml': 'yaml',
                                'yml': 'yaml'
                            }
                            
                            output_ext = ext_map.get(target_format, target_format)
                            output_filename = f"{os.path.splitext(file.name)[0]}_converted.{output_ext}"
                            
                            st.download_button(
                                label=f"Download {output_ext.upper()}",
                                data=output_data,
                                file_name=output_filename,
                                mime=mime_type,
                                key=f"dl_{file.name}"
                            )
                            
                        except Exception as e:
                            st.error(f"Conversion failed: {str(e)}")
                            logger.error(f"Conversion error for {file.name}: {str(e)}")

if __name__ == "__main__":
    main()
